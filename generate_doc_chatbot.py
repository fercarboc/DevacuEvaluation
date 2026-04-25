"""
Genera: DEBACU_Chatbot_Implantacion.docx
Guía completa de implementación del chatbot para Debacu Evaluation
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tblPr.append(borders)

def add_heading(doc, text, level=1, color_hex='1F3864'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    return p

def add_body(doc, text, bold=False, italic=False, color_hex=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    if color_hex:
        r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg='1F3864'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F2F5FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def add_phase_box(doc, phase_num, title, color_hex, items):
    p = doc.add_paragraph()
    run = p.add_run(f"  FASE {phase_num} — {title}  ")
    run.bold = True
    run.font.size = Pt(12)
    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    for item in items:
        add_bullet(doc, item)

def page_break(doc):
    doc.add_page_break()

def add_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEBACU EVALUATION")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 56, 100)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Guía Completa de Implantación del Chatbot")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(70, 130, 180)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Plan de Implementación por Fases · Arquitectura · FAQs · Entrenamiento")
    run3.font.size = Pt(12)
    run3.italic = True
    run3.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Versión 1.0 — Abril 2026")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(130, 130, 130)

    page_break(doc)

# ─── main ─────────────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Fuente por defecto
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

add_cover(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. RESUMEN EJECUTIVO", 1)
add_body(doc, (
    "Este documento define la hoja de ruta completa para implementar un chatbot conversacional "
    "en Debacu Evaluation. El chatbot tiene dos funciones principales:\n\n"
    "① Asistente de soporte para clientes (FAQs, guía de uso, resolución de dudas)\n"
    "② Asistente operativo interno (consultas de riesgo guiadas, ayuda en el panel)\n\n"
    "El plan se divide en 4 fases progresivas que permiten obtener valor desde la primera semana "
    "sin bloquear el lanzamiento por integraciones complejas."
))

add_heading(doc, "Objetivos del chatbot", 2, '2E75B6')
add_table(doc,
    ["Objetivo", "Indicador de éxito", "Fase"],
    [
        ["Reducir tickets de soporte en consultas básicas",   "−40% tickets en 60 días",     "1"],
        ["Guiar al usuario en el onboarding",                 "Perfil completado en 1ª sesión ≥70%", "2"],
        ["Responder FAQs de privacidad y RGPD",               "0 escalaciones por duda legal", "1"],
        ["Asistir en interpretación de niveles de riesgo",    "NPS chatbot ≥ 7",              "2"],
        ["Integración con funciones del panel vía API",       "Consulta de riesgo desde chat", "3"],
        ["Aprendizaje continuo y escalado a humano",          "Resolución autónoma ≥ 80%",    "4"],
    ],
    col_widths=[3.2, 2.5, 0.8]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARQUITECTURA DEL CHATBOT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. ARQUITECTURA TÉCNICA", 1)

add_heading(doc, "2.1 Capas del sistema", 2, '2E75B6')
add_table(doc,
    ["Capa", "Componente", "Tecnología recomendada", "Descripción"],
    [
        ["Interfaz", "Widget en la app", "React + iframe / Web Component", "Burbuja flotante en todas las rutas de /app/*"],
        ["Interfaz", "Web pública", "Script embebible", "Widget en /planes, /producto, /solicitar-acceso"],
        ["Orquestador", "LLM principal", "Claude API (claude-sonnet-4-6)", "Motor de razonamiento y generación de respuestas"],
        ["Conocimiento", "Base de datos RAG", "Supabase pgvector + embeddings", "FAQs, guía de uso, catálogo de incidencias"],
        ["Contexto", "Auth context", "JWT de sesión Supabase", "Plan del usuario, org_id, historial de la sesión"],
        ["Acción", "Tool calls / API", "Edge Functions de Debacu", "Consulta de riesgo, screening, exportación"],
        ["Escalado", "Handoff humano", "Email / Crisp / Intercom", "Transferencia si confianza < umbral"],
        ["Analytics", "Log de conversaciones", "Tabla chat_sessions en Supabase", "Métricas de resolución, temas recurrentes"],
    ],
    col_widths=[1.2, 1.5, 2.2, 2.5]
)

add_heading(doc, "2.2 Diagrama de flujo de una conversación", 2, '2E75B6')
add_body(doc, (
    "Usuario escribe mensaje\n"
    "        ↓\n"
    "[Clasificación de intención] → ¿es consulta de riesgo? ¿FAQ? ¿onboarding? ¿billing?\n"
    "        ↓\n"
    "[Recuperación RAG] → busca en base de conocimiento los fragmentos más relevantes\n"
    "        ↓\n"
    "[Contexto de sesión] → inyecta: plan del usuario, pantalla actual, historial reciente\n"
    "        ↓\n"
    "[LLM genera respuesta] → con los fragmentos RAG + contexto + historial\n"
    "        ↓\n"
    "¿Requiere acción en API? → Sí → llamada a Edge Function → resultado → respuesta final\n"
    "        ↓\n"
    "¿Confianza < 0.6? → Sí → ofrecer escalado a soporte humano\n"
    "        ↓\n"
    "Respuesta al usuario + log en chat_sessions"
), italic=True, color_hex='2E4057')

add_heading(doc, "2.3 Modelo de datos del chatbot (tabla a crear en Supabase)", 2, '2E75B6')
add_table(doc,
    ["Tabla", "Campo clave", "Descripción"],
    [
        ["chat_sessions",    "id, org_id, user_id, started_at, channel",          "Una sesión por conversación"],
        ["chat_messages",    "session_id, role, content, intent, confidence, ts",  "Cada mensaje entrada/salida"],
        ["chat_feedback",    "message_id, rating (1-5), comment",                  "Feedback por respuesta"],
        ["chat_kb_chunks",   "id, content, embedding, source_doc, category",       "Base de conocimiento RAG"],
        ["chat_escalations", "session_id, reason, assigned_to, resolved_at",       "Escalados a soporte humano"],
    ],
    col_widths=[2.0, 3.0, 2.4]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. FASES DE IMPLANTACIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. FASES DE IMPLANTACIÓN", 1)

# FASE 1
add_heading(doc, "FASE 1 — Chatbot estático de FAQs y soporte básico", 2, '1F5C99')
add_body(doc, "Duración estimada: 2–3 semanas | Sin integración con API | ROI inmediato")
add_body(doc, "Objetivo: Reducir el volumen de soporte básico y dar respuesta inmediata a dudas de privacidad, planes y uso general.", bold=False)

add_heading(doc, "Alcance funcional", 3, '2E75B6')
add_table(doc,
    ["Funcionalidad", "Descripción", "Fuente de datos"],
    [
        ["Responder FAQs de privacidad", "¿Están encriptados los datos? ¿Se puede identificar al huésped?", "Documento FAQ"],
        ["Explicar niveles de riesgo", "Qué significa NONE, LOW, MEDIUM, HIGH, CRITICAL", "Documento guía"],
        ["Guía de uso básica", "Cómo hacer una búsqueda, cómo registrar incidencia", "Documento guía"],
        ["Información de planes", "Diferencias BASIC/PROFESSIONAL/ENTERPRISE, precios", "Planes actuales"],
        ["Proceso de solicitud de acceso", "Pasos para darse de alta, qué documentos se necesitan", "Documento guía"],
        ["Contacto y escalado", "Redirigir a soporte si no sabe responder", "Email soporte"],
    ],
    col_widths=[2.0, 3.2, 1.5]
)

add_heading(doc, "Base de conocimiento Fase 1 — Documentos a indexar", 3, '2E75B6')
add_table(doc,
    ["Documento", "Fragmentos estimados", "Prioridad"],
    [
        ["FAQs completas (este documento)", "80–100 fragmentos", "CRÍTICA"],
        ["Guía de uso paso a paso",         "30–40 fragmentos",  "CRÍTICA"],
        ["Explicación de niveles de riesgo","10 fragmentos",      "ALTA"],
        ["Comparativa de planes",           "15 fragmentos",      "ALTA"],
        ["Política de privacidad y RGPD",   "20 fragmentos",      "ALTA"],
        ["Proceso de solicitud de acceso",  "10 fragmentos",      "MEDIA"],
        ["Glosario de términos",            "20 fragmentos",      "MEDIA"],
    ],
    col_widths=[3.0, 2.0, 1.5]
)

add_heading(doc, "Implementación técnica Fase 1", 3, '2E75B6')
add_bullet(doc, "Crear base de conocimiento en Supabase con extensión pgvector")
add_bullet(doc, "Generar embeddings de todos los fragmentos con text-embedding-3-small o equivalente")
add_bullet(doc, "Implementar Edge Function 'chatbot_query': recibe mensaje → busca top-5 fragmentos → llama a Claude → devuelve respuesta")
add_bullet(doc, "Widget React embebible en todas las páginas de /app/* con autenticación JWT")
add_bullet(doc, "Widget público (sin auth) para /planes, /producto, /solicitar-acceso")
add_bullet(doc, "Logging básico en chat_sessions + chat_messages")
add_bullet(doc, "Botón 'Hablar con soporte' siempre visible como fallback")

add_heading(doc, "Criterios de aceptación Fase 1", 3, '2E75B6')
add_bullet(doc, "El chatbot responde correctamente ≥ 85% de las preguntas del banco de pruebas")
add_bullet(doc, "Tiempo de respuesta < 3 segundos en P90")
add_bullet(doc, "Las respuestas sobre privacidad y encriptación son 100% correctas")
add_bullet(doc, "El escalado a soporte funciona en todos los casos no resueltos")

doc.add_paragraph()

# FASE 2
add_heading(doc, "FASE 2 — Chatbot contextual con awareness del usuario", 2, '1A7A3E')
add_body(doc, "Duración estimada: 2–3 semanas adicionales | Integración con contexto de sesión")
add_body(doc, "Objetivo: El chatbot sabe quién es el usuario, su plan y en qué pantalla está. Las respuestas son personalizadas.")

add_heading(doc, "Nuevas capacidades Fase 2", 3, '2E75B6')
add_table(doc,
    ["Capacidad", "Ejemplo de uso", "Datos necesarios"],
    [
        ["Conocer el plan del usuario", "'Tu plan BASIC no incluye Revenue Intelligence — ¿quieres ver cómo hacer el upgrade?'", "plan del JWT"],
        ["Saber en qué pantalla está", "Si está en /app/screening, responder con contexto de CSV", "URL actual"],
        ["Onboarding guiado", "Detectar primer login y guiar paso a paso: perfil → primera consulta", "profile_completed"],
        ["Alertas contextuales", "Si tiene alertas activas, mencionarlas proactivamente", "debacu_eval_alerts"],
        ["Historial de la sesión", "Recordar lo que preguntó antes en la misma conversación", "chat_messages"],
        ["Explicar el resultado de su última consulta", "Interpretar el risk_level que acaba de ver", "manual_check_results"],
    ],
    col_widths=[2.2, 3.2, 2.0]
)

add_heading(doc, "Implementación técnica Fase 2", 3, '2E75B6')
add_bullet(doc, "Inyectar en el prompt del sistema: plan del usuario, pantalla actual, perfil completado (sí/no), número de consultas realizadas")
add_bullet(doc, "Leer del contexto de sesión: org_id, user_id, token JWT — ya disponibles en EvalAuthContext")
add_bullet(doc, "Crear endpoint chatbot_context que devuelve el resumen de estado del usuario para inyectar al chatbot")
add_bullet(doc, "Añadir intención 'onboarding_help' que activa el flujo guiado de configuración")
add_bullet(doc, "Añadir intención 'interpretar_riesgo' que recupera el último resultado de consulta y lo explica")
add_bullet(doc, "Mejorar logging: añadir campo 'intent' y 'context_snapshot' en chat_messages")

add_heading(doc, "Criterios de aceptación Fase 2", 3, '2E75B6')
add_bullet(doc, "El chatbot menciona correctamente el plan del usuario en preguntas sobre funcionalidades")
add_bullet(doc, "El flujo de onboarding guiado se activa en el primer login en ≥ 90% de los casos")
add_bullet(doc, "La tasa de abandono del onboarding baja un 30% respecto a la línea base pre-chatbot")

doc.add_paragraph()

# FASE 3
add_heading(doc, "FASE 3 — Chatbot operativo con acciones en el sistema", 2, '7B3F00')
add_body(doc, "Duración estimada: 3–4 semanas adicionales | Tool calls a Edge Functions")
add_body(doc, "Objetivo: El chatbot no solo informa, sino que puede ejecutar acciones en nombre del usuario.")

add_heading(doc, "Acciones disponibles en Fase 3", 3, '2E75B6')
add_table(doc,
    ["Acción", "Edge Function invocada", "Requiere confirmación del usuario"],
    [
        ["Hacer una consulta de riesgo",         "debacu_eval_manual_check",              "No (acción de solo lectura)"],
        ["Ver mis últimas consultas",            "debacu_eval_manual_check_mine",          "No"],
        ["Ver mis alertas activas",              "debacu_eval_get_risk_alerts",            "No"],
        ["Ver el estado de mi suscripción",      "debacu_eval_subscription_state_get",     "No"],
        ["Iniciar proceso de cambio de plan",    "debacu_eval_subscription_change_plan",   "Sí — confirmación explícita"],
        ["Ver mi historial de exportaciones",    "debacu_eval_audit_history_list",         "No"],
        ["Descargar última exportación",         "customer_audit_export_download",         "Sí"],
        ["Comprobar estado de integración PMS",  "pms-connection-test",                    "No"],
    ],
    col_widths=[2.5, 2.8, 2.0]
)

add_heading(doc, "Protocolo de seguridad para acciones", 3, '2E75B6')
add_bullet(doc, "Todas las acciones se ejecutan con el JWT del usuario — nunca con credenciales elevadas")
add_bullet(doc, "Acciones de escritura (cambio de plan, exportación) requieren confirmación explícita del usuario ('¿confirmas que quieres...?')")
add_bullet(doc, "El chatbot nunca puede registrar incidencias automáticamente — solo puede guiar al usuario para que lo haga él")
add_bullet(doc, "Si una action falla, el chatbot informa del error y ofrece ir directamente a la pantalla correspondiente")
add_bullet(doc, "Todas las acciones ejecutadas via chatbot se registran en debacu_eval_audit_log con source='chatbot'")

add_heading(doc, "Implementación técnica Fase 3", 3, '2E75B6')
add_bullet(doc, "Definir tools en el prompt del sistema de Claude: cada tool con nombre, descripción y parámetros")
add_bullet(doc, "Crear middleware chatbot_action_handler que valida el JWT antes de ejecutar cualquier tool call")
add_bullet(doc, "Implementar función de routing: mensaje del chatbot con tool_use → Edge Function correspondiente → resultado → respuesta natural")
add_bullet(doc, "Añadir campo 'action_taken' y 'action_result' en chat_messages para auditoría")

add_heading(doc, "Criterios de aceptación Fase 3", 3, '2E75B6')
add_bullet(doc, "Las consultas de riesgo via chatbot devuelven resultados idénticos a los de la pantalla de búsqueda")
add_bullet(doc, "Ninguna acción de escritura se ejecuta sin confirmación explícita del usuario")
add_bullet(doc, "El 100% de las acciones via chatbot aparecen en el audit_log")

doc.add_paragraph()

# FASE 4
add_heading(doc, "FASE 4 — Chatbot inteligente con aprendizaje y métricas", 2, '5C0099')
add_body(doc, "Duración estimada: continuo | Mejora iterativa basada en datos reales")
add_body(doc, "Objetivo: El chatbot mejora con el uso, detecta patrones, reduce la tasa de escalado y alimenta el roadmap del producto.")

add_heading(doc, "Capacidades Fase 4", 3, '2E75B6')
add_table(doc,
    ["Capacidad", "Descripción", "Valor de negocio"],
    [
        ["Análisis de intenciones no cubiertas", "Detectar preguntas frecuentes sin respuesta buena", "Alimenta roadmap de producto"],
        ["Fine-tuning de la KB",                 "Añadir nuevos fragmentos basados en preguntas reales", "Mejora tasa de resolución"],
        ["Sistema de feedback por mensaje",      "Rating 👍/👎 en cada respuesta + comentario opcional", "Identifica respuestas malas"],
        ["Detección de frustración",             "Si el usuario reformula 3+ veces, ofrecer soporte proactivamente", "Reduce abandono"],
        ["Resúmenes de uso para el admin",       "Dashboard de temas más consultados, tasa de resolución, NPS", "ROI visible"],
        ["Personalización por tipo de alojamiento","Respuestas distintas para hotel vs apartamento vs camping", "Relevancia aumentada"],
        ["Proactive messaging",                  "Si el sistema detecta que el perfil no está completo, el chatbot aparece proactivamente", "Mejora activación"],
    ],
    col_widths=[2.2, 2.8, 2.5]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. BASE DE CONOCIMIENTO COMPLETA — FAQs
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. BASE DE CONOCIMIENTO — FAQs COMPLETAS PARA EL CHATBOT", 1)
add_body(doc, (
    "Esta sección contiene el banco de preguntas y respuestas que debe cargarse como base de conocimiento RAG. "
    "Cada par pregunta-respuesta debe indexarse como un fragmento independiente con sus variantes de pregunta."
))

# ── Bloque: Producto
add_heading(doc, "4.1 Sobre el producto", 2, '2E75B6')

faqs_producto = [
    ("¿Qué es Debacu Evaluation?",
     "Debacu Evaluation es una plataforma online para hoteles y alojamientos que permite consultar el historial de riesgo de un huésped antes de que llegue, y registrar incidencias cuando ocurren. Todo el proceso está diseñado para cumplir con el RGPD: los datos están hasheados de forma irreversible y nunca se almacena ningún dato personal identificable en claro."),
    ("¿Para qué tipo de negocio es útil Debacu Evaluation?",
     "Para cualquier negocio de alojamiento con riesgo asociado a clientes: hoteles urbanos y vacacionales, apartamentos turísticos de corta, media y larga estancia, casas rurales, hostales, posadas, campings y cualquier establecimiento que quiera gestionar el riesgo de impago, daños o comportamientos problemáticos."),
    ("¿Cómo me ayuda en mi hotel?",
     "Te permite saber, antes del check-in, si un huésped tiene historial de problemas en otros establecimientos. También te ayuda a documentar incidencias con valor económico, analizar tus ingresos por canal, y detectar anomalías en tus reservas. En definitiva, reduces pérdidas y tomas mejores decisiones."),
    ("¿Necesito conocimientos técnicos para usar la plataforma?",
     "No. La plataforma está diseñada para equipos de recepción y dirección sin perfil técnico. Si puedes usar una hoja de cálculo, puedes usar Debacu Evaluation."),
]

for q, a in faqs_producto:
    add_body(doc, f"P: {q}", bold=True)
    add_body(doc, f"R: {a}")
    doc.add_paragraph()

# ── Bloque: Privacidad
add_heading(doc, "4.2 Sobre privacidad y datos personales", 2, '2E75B6')

faqs_privacidad = [
    ("¿Los datos de los huéspedes están encriptados?",
     "Sí. Todos los datos sensibles están encriptados. Además, los identificadores (documento, email, teléfono) se transforman en hashes criptográficos irreversibles antes de almacenarse. Es imposible reconstruir el dato original a partir de lo que hay en la base de datos."),
    ("¿Se puede identificar directamente a un huésped?",
     "No. Debacu Evaluation no almacena nombres, documentos ni emails en claro en su base de datos compartida. Solo almacena hashes irreversibles. Ni el equipo de Debacu ni ningún otro establecimiento puede saber la identidad de un huésped a partir de los datos del sistema."),
    ("¿Cumple con el RGPD?",
     "Sí. El sistema incluye: firma digital del Acuerdo de Protección de Datos (DPA), anonimización irreversible de todos los identificadores, auditoría completa de todos los accesos y exportaciones, y registro de base legal para cada tratamiento de datos. Debacu actúa como encargado del tratamiento bajo tu responsabilidad como responsable."),
    ("¿Quién puede ver los datos que introduzco?",
     "Solo tu organización. El sistema tiene aislamiento completo entre clientes mediante Row Level Security. El equipo técnico de Debacu solo puede acceder en caso de soporte crítico, y ese acceso queda registrado."),
    ("¿Los datos que registro los pueden ver otros hoteles?",
     "No en detalle. Otros establecimientos solo pueden saber el nivel de riesgo agregado de un huésped (NONE/LOW/MEDIUM/HIGH/CRITICAL). Nunca pueden saber que fue tu establecimiento quien registró la incidencia, ni ver el contenido de lo que registraste."),
    ("¿Puedo exportar mis datos para dárselos a mi abogado o a la policía?",
     "Sí. En 'Auditoría → Exportaciones' puedes generar un informe firmado con hash SHA-256 que garantiza su integridad legal. Puedes indicar el destinatario y la base legal de la entrega. El sistema registra automáticamente cada exportación."),
    ("¿Qué pasa si un huésped solicita que elimine sus datos?",
     "Debes contactar con el soporte de Debacu indicando el ejercicio del derecho de supresión. El proceso incluye la eliminación del identity_key asociado conforme al RGPD. Te proporcionaremos la documentación necesaria para responder al interesado."),
    ("¿Puedo usar estos datos en un juicio?",
     "Las exportaciones de Debacu incluyen hash de integridad del archivo (SHA-256), fecha y hora de generación, y registro del destinatario. Esto les da valor como evidencia documental. Para su uso en procedimientos legales, te recomendamos consultarlo con tu asesor jurídico."),
]

for q, a in faqs_privacidad:
    add_body(doc, f"P: {q}", bold=True)
    add_body(doc, f"R: {a}")
    doc.add_paragraph()

# ── Bloque: Uso
add_heading(doc, "4.3 Sobre el uso del sistema", 2, '2E75B6')

faqs_uso = [
    ("¿Cómo hago una consulta de riesgo?",
     "Ve al menú lateral y pulsa 'Búsqueda' (o 'Consultar huésped'). Introduce el documento, email o teléfono del huésped y pulsa Consultar. En menos de 2 segundos recibes el nivel de riesgo con los datos disponibles."),
    ("¿Con qué datos puedo consultar?",
     "Con el número de documento de identidad, el email o el teléfono. Puedes usar uno solo o varios a la vez. Cuantos más datos aportes, más precisa es la identificación."),
    ("¿Qué significa cada nivel de riesgo?",
     "SIN HISTORIAL: no aparece en el sistema, procede con normalidad. BAJO: algún incidente aislado de poco impacto. MEDIO: varios incidentes o impacto económico notable, considera pedir depósito. ALTO: historial serio con pérdidas significativas, máximas garantías o reconsiderar. CRÍTICO: múltiples incidentes graves, recomendable denegar el servicio."),
    ("¿Cómo añado un cliente conflictivo?",
     "Ve a 'Registrar incidencia' en el menú. Introduce los datos del huésped (documento, email o teléfono), selecciona el tipo de incidencia del catálogo, indica el importe económico si aplica, añade notas y guarda. El sistema actualizará automáticamente el índice de riesgo de ese huésped."),
    ("¿Puedo registrar una incidencia de un huésped que ya se fue?",
     "Sí. No hay límite de tiempo para registrar incidencias pasadas. Solo necesitas tener los datos identificativos del huésped."),
    ("¿Qué pasa si me equivoco al introducir datos?",
     "Puedes editar o eliminar una incidencia que hayas registrado tú mismo. El sistema registra la modificación en el historial de auditoría. Para incidencias de otros establecimientos, contacta con soporte."),
    ("¿Cómo veo el riesgo en un período concreto?",
     "En 'Revenue → Riesgo' puedes seleccionar un rango de fechas y ver qué reservas del período tienen huéspedes con historial. También en 'Auditoría → Histórico' puedes filtrar consultas por fecha."),
    ("¿Cómo funciona el Screening CSV?",
     "Sube un CSV con tu lista de huéspedes (plantilla descargable desde la pantalla de Screening). El sistema analiza cada huésped y te devuelve un informe con cuántos tienen riesgo alto, medio y bajo. Puedes hacerlo antes de que lleguen para prepararte."),
    ("¿Cómo veo el histórico de consultas que he hecho?",
     "En 'Auditoría → Histórico'. Aparecen todas las consultas con fecha, tipo de búsqueda y resultado enmascarado."),
]

for q, a in faqs_uso:
    add_body(doc, f"P: {q}", bold=True)
    add_body(doc, f"R: {a}")
    doc.add_paragraph()

# ── Bloque: Planes
add_heading(doc, "4.4 Sobre planes y pagos", 2, '2E75B6')

faqs_planes = [
    ("¿Cuáles son los planes disponibles?",
     "BASIC: consulta manual de riesgo, registro de incidencias, auditoría básica y reporte semanal PDF. PROFESSIONAL: todo lo de BASIC más Revenue Intelligence completo (canales, pickup, pricing, temporadas). ENTERPRISE: todo lo de PROFESSIONAL más integración con tu PMS, exportaciones avanzadas y gestión multi-propiedad."),
    ("¿Cómo cambio mi plan?",
     "En 'Mi cuenta → Plan y suscripción'. Pulsa 'Cambiar plan', selecciona el nuevo y confirma. El upgrade es inmediato. El downgrade se aplica al final del período de facturación actual."),
    ("¿Qué pierdo si bajo de plan?",
     "Si bajas de PROFESSIONAL a BASIC, perderás acceso a Revenue Intelligence y sus datos históricos quedarán en pausa hasta que vuelvas a subir. Si bajas de ENTERPRISE a PROFESSIONAL, perderás la integración con el PMS. Los datos no se eliminan."),
    ("¿Hay período de prueba?",
     "Sí, existe un período de prueba al dar de alta la cuenta. Consulta las condiciones actuales al solicitar acceso o pregunta al equipo de soporte."),
    ("¿Cómo pago?",
     "El pago se gestiona mediante Stripe, con tarjeta de crédito/débito. También puedes configurar domiciliación bancaria SEPA. Tus datos de pago están protegidos por Stripe y no los almacenamos en nuestros servidores."),
    ("¿Dónde están mis facturas?",
     "En 'Mi cuenta → Facturación'. Puedes ver y descargar todas las facturas generadas."),
    ("¿Qué pasa si no pago una factura?",
     "Si el pago falla, recibirás un email de aviso. Tienes un período de gracia para regularizarlo. Si no se regulariza, la cuenta pasa a estado suspendido y no podrás hacer consultas hasta que el pago esté al día."),
]

for q, a in faqs_planes:
    add_body(doc, f"P: {q}", bold=True)
    add_body(doc, f"R: {a}")
    doc.add_paragraph()

# ── Bloque: Revenue
add_heading(doc, "4.5 Sobre Revenue Intelligence", 2, '2E75B6')

faqs_revenue = [
    ("¿Qué es Revenue Intelligence?",
     "Es el módulo de análisis de ingresos incluido en planes PROFESSIONAL y ENTERPRISE. Te permite ver de dónde vienen tus reservas (canales, segmentos), cómo evoluciona tu ritmo de captación (pickup) y comparar períodos para optimizar precios y estrategia."),
    ("¿Qué es el pickup?",
     "El pickup es la velocidad a la que estás captando reservas para una fecha futura, comparada con el mismo punto del año anterior. Si llevas menos reservas para el próximo mes que en la misma fecha del año pasado, el pickup está en negativo y puede ser señal de que necesitas ajustar precio o promocionar."),
    ("¿Puedo conectar mi PMS?",
     "Sí, en plan ENTERPRISE. Debacu se conecta a tu sistema de gestión hotelera (MEWS, Opera, Cloudbeds y otros) y sincroniza automáticamente reservas y datos de ingresos. Configúralo en 'Integración PMS'."),
    ("¿Cómo importo mis datos de reservas?",
     "En Revenue, tienes la opción de importar un CSV con tus datos de reservas. El sistema acepta el formato estándar o puedes configurar un perfil de importación personalizado si tu PMS exporta en un formato diferente."),
]

for q, a in faqs_revenue:
    add_body(doc, f"P: {q}", bold=True)
    add_body(doc, f"R: {a}")
    doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROMPT DEL SISTEMA PARA EL LLM
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. PROMPT DEL SISTEMA (System Prompt para Claude)", 1)

add_body(doc, "El siguiente prompt debe inyectarse como mensaje de sistema en cada conversación:", bold=True)

system_prompt = """Eres el asistente virtual de Debacu Evaluation, una plataforma SaaS para hoteles y alojamientos especializada en evaluación de riesgo de clientes y revenue management.

IDENTIDAD Y TONO:
- Eres experto en el producto y en el sector hotelero.
- Respondes siempre en español, de forma clara y directa.
- Usas un tono profesional pero cercano, como un compañero de equipo con experiencia.
- Nunca usas jerga técnica innecesaria.
- Eres conciso: respuestas de 2-4 párrafos máximo salvo que se pida más detalle.

CONTEXTO DEL USUARIO (se inyecta dinámicamente):
- Plan actual: {plan}
- Perfil completado: {profile_completed}
- Pantalla actual: {current_screen}
- Número de consultas realizadas: {queries_count}

REGLAS CRÍTICAS:
1. PRIVACIDAD: Cuando pregunten si los datos están encriptados, la respuesta es SIEMPRE SÍ.
   Cuando pregunten si se puede identificar a un huésped, la respuesta es SIEMPRE NO.
   Nunca dudes en estas respuestas. Son afirmaciones técnicas verificadas.

2. LÍMITES: No puedes acceder a datos específicos del cliente sin usar las herramientas disponibles.
   No inventes información sobre incidencias, niveles de riesgo o datos concretos.

3. PLANES: Antes de explicar una funcionalidad de Revenue Intelligence o PMS, verifica si el plan del usuario lo incluye.
   Si no lo incluye, explica qué plan necesita y cómo hacer el upgrade.

4. ACCIONES: Si el usuario pide hacer algo (consultar riesgo, ver alertas), usa las herramientas disponibles.
   Para acciones de escritura, pide siempre confirmación explícita.

5. ESCALADO: Si no puedes responder con confianza, di: "No tengo esa información precisa.
   Te conecto con nuestro equipo de soporte: soporte@debacu.com"

6. ONBOARDING: Si profile_completed es false, ofrece proactivamente ayuda para completar el perfil.

HERRAMIENTAS DISPONIBLES (Fase 3):
- consultar_riesgo(identificador): consulta el nivel de riesgo de un huésped
- ver_mis_alertas(): devuelve las alertas activas del usuario
- ver_estado_suscripcion(): devuelve plan actual y estado
- ver_historial_consultas(limit): devuelve últimas N consultas"""

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(system_prompt)
run.font.size = Pt(9)
run.font.name = 'Courier New'
run.font.color.rgb = RGBColor(30, 60, 30)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. GLOSARIO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. GLOSARIO DE TÉRMINOS PARA LA BASE DE CONOCIMIENTO", 1)

add_table(doc,
    ["Término", "Definición para el chatbot"],
    [
        ["identity_key",      "Hash criptográfico único que identifica a un huésped sin almacenar sus datos personales"],
        ["risk_band",         "Clasificación de riesgo: HIGH (alto), MEDIUM (medio), LOW (bajo), NONE (sin historial)"],
        ["risk_level",        "Clasificación detallada: CRITICAL, HIGH, MEDIUM, LOW, NONE"],
        ["risk_score",        "Puntuación numérica de 0 a 100 que refleja el nivel de riesgo calculado"],
        ["screening CSV",     "Proceso de análisis masivo de una lista de huéspedes cargada desde un archivo CSV"],
        ["import batch",      "Lote de importación: un grupo de registros subidos en una sola operación"],
        ["org_id",            "Identificador único de tu organización (tu establecimiento o cadena)"],
        ["DPA",               "Acuerdo de Protección de Datos — contrato firmado digitalmente al darte de alta"],
        ["pickup",            "Velocidad de captación de reservas para una fecha futura vs el mismo punto del año anterior"],
        ["ADR",               "Average Daily Rate — precio medio por habitación ocupada"],
        ["RevPAR",            "Revenue Per Available Room — ingreso por habitación disponible (ocupadas y libres)"],
        ["PMS",               "Property Management System — sistema de gestión hotelera (MEWS, Opera, Cloudbeds, etc.)"],
        ["RLS",               "Row Level Security — mecanismo de Supabase que aísla los datos entre organizaciones"],
        ["JWT",               "JSON Web Token — credencial cifrada que autentica al usuario en cada petición"],
        ["pepper",            "Clave secreta global que se combina con los datos al generar el hash, haciendo irreversible la identificación"],
        ["plan BASIC",        "Plan básico: consulta manual, registro de incidencias, auditoría y reporte semanal"],
        ["plan PROFESSIONAL", "Plan medio: todo lo de BASIC más Revenue Intelligence completo"],
        ["plan ENTERPRISE",   "Plan avanzado: todo lo de PROFESSIONAL más integración PMS y multi-propiedad"],
        ["profile_completed", "Indicador de si el cliente ha completado el perfil de su establecimiento"],
        ["SHA-256",           "Algoritmo de hash usado para verificar la integridad de las exportaciones"],
    ],
    col_widths=[2.0, 5.4]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. CHECKLIST DE IMPLANTACIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. CHECKLIST DE IMPLANTACIÓN POR FASE", 1)

add_heading(doc, "Fase 1 — Checklist", 2, '1F5C99')
f1_items = [
    "[ ] Crear tabla chat_kb_chunks con columna embedding (pgvector)",
    "[ ] Generar embeddings de todos los fragmentos de FAQ y guía de uso",
    "[ ] Implementar Edge Function chatbot_query (RAG + Claude)",
    "[ ] Crear componente React ChatWidget (burbuja flotante)",
    "[ ] Integrar widget en rutas /app/* con auth JWT",
    "[ ] Integrar widget público en /planes y /solicitar-acceso",
    "[ ] Crear tablas chat_sessions y chat_messages",
    "[ ] Implementar botón de escalado a soporte (email)",
    "[ ] Tests de regresión: 50 preguntas del banco de pruebas",
    "[ ] Desplegar en staging y validar con equipo interno",
    "[ ] Go-live Fase 1",
]
for item in f1_items:
    add_bullet(doc, item)

add_heading(doc, "Fase 2 — Checklist", 2, '1A7A3E')
f2_items = [
    "[ ] Implementar Edge Function chatbot_context (devuelve estado del usuario)",
    "[ ] Modificar system prompt para incluir contexto dinámico del usuario",
    "[ ] Implementar detección de intención 'onboarding_help'",
    "[ ] Implementar flujo guiado de onboarding en el chatbot",
    "[ ] Añadir campo 'intent' en chat_messages",
    "[ ] Tests A/B: onboarding con y sin chatbot",
    "[ ] Monitoreo de tasa de completado de perfil",
]
for item in f2_items:
    add_bullet(doc, item)

add_heading(doc, "Fase 3 — Checklist", 2, '7B3F00')
f3_items = [
    "[ ] Definir tools en el system prompt de Claude",
    "[ ] Implementar middleware chatbot_action_handler con validación JWT",
    "[ ] Conectar tool consultar_riesgo → debacu_eval_manual_check",
    "[ ] Conectar tool ver_mis_alertas → debacu_eval_get_risk_alerts",
    "[ ] Conectar tool ver_estado_suscripcion → debacu_eval_subscription_state_get",
    "[ ] Implementar flujo de confirmación para acciones de escritura",
    "[ ] Añadir source='chatbot' en todas las entradas de audit_log",
    "[ ] Tests de seguridad: intentos de acceso cross-org",
    "[ ] Revisión legal de acciones permitidas via chatbot",
]
for item in f3_items:
    add_bullet(doc, item)

add_heading(doc, "Fase 4 — Checklist", 2, '5C0099')
f4_items = [
    "[ ] Implementar sistema de feedback 👍/👎 por mensaje",
    "[ ] Crear tabla chat_feedback",
    "[ ] Dashboard de métricas del chatbot (temas, resolución, NPS)",
    "[ ] Proceso mensual de revisión y enriquecimiento de la KB",
    "[ ] Implementar detección de frustración (≥3 reformulaciones)",
    "[ ] Segmentación de respuestas por tipo de alojamiento",
    "[ ] Revisar y actualizar este documento trimestralmente",
]
for item in f4_items:
    add_bullet(doc, item)

# ── Guardar
output_path = r"c:\DebacuEvaluation\DEBACU_Chatbot_Implantacion.docx"
doc.save(output_path)
print(f"Documento guardado: {output_path}")
