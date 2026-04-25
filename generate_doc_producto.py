"""
Genera: DEBACU_Producto_Estado_y_Plan_de_Accion.docx
Análisis del estado actual + necesidades + plan de acción por fases
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tblPr.append(borders)

def add_heading(doc, text, level=1, color_hex='1F3864'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    return p

def add_body(doc, text, bold=False, italic=False, color_hex=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    if color_hex:
        r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg='1F3864'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F2F5FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def page_break(doc):
    doc.add_page_break()

def add_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEBACU EVALUATION")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 56, 100)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Estado del Producto, Análisis de Necesidades")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(70, 130, 180)

    p2b = doc.add_paragraph()
    p2b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2b = p2b.add_run("y Plan de Acción por Fases")
    run2b.font.size = Pt(18)
    run2b.font.color.rgb = RGBColor(70, 130, 180)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Análisis interno · Base de datos · Roadmap de mejoras · Arquitectura")
    run3.font.size = Pt(12)
    run3.italic = True
    run3.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Versión 1.0 — Abril 2026")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(130, 130, 130)

    page_break(doc)

# ─── main ─────────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

add_cover(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. RESUMEN EJECUTIVO", 1)
add_body(doc, (
    "Debacu Evaluation es una plataforma SaaS B2B para evaluación de riesgo de clientes en el sector hospitalidad "
    "y gestión de revenue, construida sobre Supabase (PostgreSQL + Edge Functions), React 19 y TypeScript. "
    "El sistema está en producción con un núcleo funcional sólido, pero presenta deuda técnica acumulada, "
    "funcionalidades de alto valor sin exponer en la UI, y gaps de UX que afectan directamente a la retención y activación.\n\n"
    "Este documento analiza el estado actual del producto con profundidad técnica y de negocio, identifica las "
    "necesidades prioritarias y define un plan de acción estructurado en 4 fases para convertir Debacu Evaluation "
    "en el producto de referencia en su nicho."
))

add_heading(doc, "Métricas del estado actual del sistema", 2, '2E75B6')
add_table(doc,
    ["Dimensión", "Estado actual", "Valoración"],
    [
        ["Funcionalidad core (riesgo)",      "Implementada y en producción",                           "BUENO"],
        ["Funcionalidad revenue",            "Implementada, acceso por plan",                          "BUENO"],
        ["Integración PMS",                  "Implementada (ENTERPRISE), multiples PMS",               "BUENO"],
        ["UX/UI onboarding",                 "Sin flujo guiado en primer login",                       "MEJORABLE"],
        ["Notificaciones proactivas",         "No implementadas",                                       "GAP CRÍTICO"],
        ["Chatbot / asistente",              "No existe",                                               "GAP CRÍTICO"],
        ["Watchlist UI",                     "Tabla en BD existe, sin UI",                             "GAP"],
        ["API pública documentada",          "Token existe en BD, sin documentación",                  "GAP"],
        ["Deuda técnica BD",                 "157 Edge Functions, tablas backup en producción",         "ATENCIÓN"],
        ["Seguridad y privacidad",           "Excelente (RLS, hashing, DPA, auditoría)",               "EXCELENTE"],
        ["Modelo multi-tenant",              "Robusto, RLS en todas las tablas operativas",            "EXCELENTE"],
        ["Facturación",                      "Stripe integrado, webhooks activos",                     "BUENO"],
    ],
    col_widths=[2.5, 3.5, 1.5]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ESTADO ACTUAL DEL PRODUCTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. ESTADO ACTUAL DEL PRODUCTO", 1)

add_heading(doc, "2.1 Stack tecnológico", 2, '2E75B6')
add_table(doc,
    ["Capa", "Tecnología", "Versión", "Estado"],
    [
        ["Frontend",         "React + TypeScript",         "19.2.1 / 5.8.2",   "Producción"],
        ["Router",           "React Router DOM",           "7.12.0",            "Producción"],
        ["UI Components",    "Material-UI + Tailwind",     "7.3.7 / 4.1.18",   "Producción"],
        ["Build",            "Vite",                       "6.2.0",             "Producción"],
        ["Backend/API",      "Supabase Edge Functions",    "Deno runtime",      "157 funciones activas"],
        ["Base de datos",    "PostgreSQL via Supabase",    "13.0.5",            "Producción"],
        ["Auth",             "Supabase Auth + JWT custom", "—",                 "Producción"],
        ["Pagos",            "Stripe",                     "—",                 "Producción"],
        ["Email",            "Brevo (Sendinblue)",         "—",                 "Producción"],
        ["Gráficos",         "Recharts",                   "—",                 "Producción"],
        ["Animaciones",      "Framer Motion",              "—",                 "Producción"],
    ],
    col_widths=[2.0, 2.5, 1.8, 2.0]
)

add_heading(doc, "2.2 Pantallas del sistema — inventario completo", 2, '2E75B6')
add_table(doc,
    ["Área", "Pantalla", "Ruta", "Plan mínimo"],
    [
        ["Pública",     "Landing page",               "/",                           "—"],
        ["Pública",     "Planes",                     "/planes",                     "—"],
        ["Pública",     "Solicitar acceso",           "/solicitar-acceso",           "—"],
        ["Auth",        "Login",                      "/login",                      "—"],
        ["Auth",        "Activación de cuenta",       "/auth/activate",              "—"],
        ["Auth",        "Reset contraseña",           "/auth/reset",                 "—"],
        ["App",         "Dashboard",                  "/app/dashboard",              "BASIC"],
        ["App",         "Búsqueda / Consulta riesgo", "/app/buscar",                 "BASIC"],
        ["App",         "Registrar incidencia",       "/app/registrar",              "BASIC"],
        ["App",         "Alertas",                    "/app/alarmas",                "BASIC"],
        ["App",         "Screening CSV",              "/app/screening",              "BASIC"],
        ["App",         "Revenue - Canales",          "/app/revenue/canales",        "PROFESSIONAL"],
        ["App",         "Revenue - Riesgo",           "/app/revenue/riesgo",         "PROFESSIONAL"],
        ["App",         "Revenue - Día x Día",        "/app/revenue/dia-x-dia",      "PROFESSIONAL"],
        ["App",         "Revenue - Mensual",          "/app/revenue/mensual",        "PROFESSIONAL"],
        ["App",         "Revenue - Tipos hab.",       "/app/revenue/tipos-hab",      "PROFESSIONAL"],
        ["App",         "Revenue - Pricing",          "/app/revenue/pricing",        "PROFESSIONAL"],
        ["App",         "Revenue - Temporadas",       "/app/revenue/temporadas",     "PROFESSIONAL"],
        ["App",         "Revenue - Propiedades",      "/app/revenue/propiedades",    "PROFESSIONAL"],
        ["App",         "Auditoría - Estadísticas",   "/app/auditoria/estadisticas", "BASIC"],
        ["App",         "Auditoría - Histórico",      "/app/auditoria/historico",    "BASIC"],
        ["App",         "Auditoría - Exportaciones",  "/app/auditoria/exportaciones","BASIC"],
        ["App",         "Integración PMS",            "/app/pms",                    "ENTERPRISE"],
        ["App",         "Mi cuenta",                  "/app/cuenta",                 "BASIC"],
        ["Admin",       "Dashboard admin",            "/app/admin/dashboard",        "ADMIN"],
        ["Admin",       "Solicitudes de acceso",      "/app/admin/solicitudes",      "ADMIN"],
        ["Admin",       "Clientes",                   "/app/admin/clientes",         "ADMIN"],
        ["Admin",       "Planes",                     "/app/admin/planes",           "ADMIN"],
        ["Admin",       "Facturación",                "/app/admin/facturacion",      "ADMIN"],
        ["Admin",       "Abusos",                     "/app/admin/abusos",           "ADMIN"],
        ["Admin",       "Estadísticas",               "/app/admin/estadisticas",     "ADMIN"],
        ["Admin",       "Auditoría",                  "/app/admin/auditoria",        "ADMIN"],
        ["Admin",       "Exportaciones",              "/app/admin/exportaciones",    "ADMIN"],
        ["Admin",       "Cambios de config",          "/app/admin/cambios",          "ADMIN"],
        ["Admin",       "Configuración",              "/app/admin/configuracion",    "ADMIN"],
    ],
    col_widths=[1.3, 2.5, 3.0, 1.5]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. ANÁLISIS DE BASE DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. ESQUEMA DE BASE DE DATOS — ANÁLISIS COMPLETO", 1)

add_body(doc, (
    "El proyecto cuenta con más de 140 tablas/vistas en el schema público de Supabase. "
    "A continuación se documentan las tablas clave por dominio, con sus campos principales y relaciones."
))

add_heading(doc, "3.1 Tablas por dominio", 2, '2E75B6')

add_heading(doc, "Dominio: Clientes y Organizaciones", 3, '2E75B6')
add_table(doc,
    ["Tabla", "Descripción", "Tipo"],
    [
        ["customers",                        "Registro principal de cada cliente de la plataforma",       "Núcleo"],
        ["debacu_eval_organizations",        "Organización multi-tenant (1:1 con customer en producción)","Núcleo"],
        ["debacu_eval_org_members",          "Usuarios miembros de cada organización",                    "Núcleo"],
        ["debacu_eval_hotel_profile",        "Perfil detallado del establecimiento (ADR, targets, etc.)","Núcleo"],
        ["debacu_eval_customer_profile",     "Perfil básico del cliente (tipo, habitaciones, web)",       "Auxiliar"],
        ["debacu_eval_customer_org_map",     "Mapeo customer_id ↔ org_id",                               "Auxiliar"],
        ["debacu_eval_access_requests",      "Solicitudes de acceso pendientes de aprobación",            "Núcleo"],
        ["debacu_eval_org_member_profiles",  "Perfiles extendidos de miembros",                          "Auxiliar"],
    ],
    col_widths=[3.0, 3.5, 1.2]
)

add_heading(doc, "Dominio: Riesgo e Identidad (núcleo del negocio)", 3, '2E75B6')
add_table(doc,
    ["Tabla", "Descripción", "Tipo"],
    [
        ["debacu_eval_guest_index",           "Índice global de huéspedes con riesgo acumulado (solo hashes)", "CRÍTICA"],
        ["debacu_eval_identity_risk_state",   "Estado de riesgo calculado por identity_key",                   "CRÍTICA"],
        ["debacu_eval_identity_risk_events",  "Log inmutable de todos los cambios de riesgo",                   "CRÍTICA"],
        ["debacu_eval_manual_checks",         "Log de cada consulta de riesgo realizada",                       "Núcleo"],
        ["debacu_eval_manual_check_results",  "Resultado de cada consulta (datos enmascarados)",                "Núcleo"],
        ["debacu_eval_manual_incidents",      "Incidencias registradas manualmente",                            "Núcleo"],
        ["debacu_eval_risk_alerts",           "Alertas de riesgo generadas automáticamente",                    "Núcleo"],
        ["debacu_eval_org_guest_index",       "Índice de huéspedes por organización",                          "Auxiliar"],
        ["debacu_eval_org_guest_evidence",    "Evidencias de incidencias por organización",                     "Auxiliar"],
        ["watchlist_reservations",            "Huéspedes en vigilancia activa (sin UI actualmente)",            "Auxiliar"],
    ],
    col_widths=[3.0, 3.5, 1.2]
)

add_heading(doc, "Dominio: Importación y Screening", 3, '2E75B6')
add_table(doc,
    ["Tabla", "Descripción"],
    [
        ["debacu_eval_unified_import_batches", "Lotes de importación multi-fuente (sistema unificado actual)"],
        ["debacu_eval_unified_import_rows",    "Filas individuales del sistema unificado"],
        ["debacu_eval_import_batches",         "Lotes de importación (sistema anterior, aún activo)"],
        ["debacu_eval_import_rows",            "Filas del sistema anterior"],
        ["screening_runs",                     "Ejecuciones de screening CSV"],
        ["screening_results",                  "Resultados por fila de cada screening run"],
        ["screening_alerts",                   "Alertas generadas por screening"],
        ["import_profiles",                    "Perfiles de mapeo para CSVs sin cabeceras estándar"],
        ["import_jobs",                        "Jobs de importación (legacy)"],
    ],
    col_widths=[3.0, 4.4]
)

add_heading(doc, "Dominio: Reservas y Revenue", 3, '2E75B6')
add_table(doc,
    ["Tabla", "Descripción"],
    [
        ["debacu_eval_reservations",           "Reservas consolidadas (versión viva, upsert por reservation_key)"],
        ["debacu_eval_reservation_snapshots",  "Snapshots históricos de cada reserva"],
        ["debacu_eval_reservation_identities", "Identidades de reservas: hashes + nombre encriptado"],
        ["debacu_eval_guest_stays",            "Estancias de huéspedes (checkin/checkout + identity_key)"],
        ["debacu_eval_revenue_daily",          "Revenue agregado por día y propiedad"],
        ["debacu_eval_revenue_booking_lines",  "Líneas de reserva para análisis detallado"],
        ["debacu_eval_revenue_pickup_snapshots","Snapshots de pickup para comparativas"],
        ["debacu_eval_inventory_daily",        "Habitaciones disponibles por día"],
        ["debacu_eval_room_prices",            "Precios por tipo de habitación y fecha"],
        ["debacu_eval_property_seasons",       "Temporadas y eventos con ajustes de precio"],
    ],
    col_widths=[3.0, 4.4]
)

add_heading(doc, "Dominio: Suscripciones y Facturación", 3, '2E75B6')
add_table(doc,
    ["Tabla", "Descripción"],
    [
        ["subscriptions",        "Suscripciones activas e históricas (Stripe sync)"],
        ["plans",                "Definición de planes: BASIC/MEDIUM(PROFESSIONAL)/PREMIUM(ENTERPRISE)"],
        ["debacu_eval_invoices", "Facturas sincronizadas desde Stripe"],
        ["debacu_eval_payments", "Pagos registrados"],
        ["subscription_events",  "Log de upgrades, downgrades, cancelaciones"],
    ],
    col_widths=[2.5, 4.9]
)

add_heading(doc, "3.2 Campos clave de las tablas más importantes", 2, '2E75B6')

add_heading(doc, "customers — tabla raíz de cada cliente", 3, '2E75B6')
add_table(doc,
    ["Campo", "Tipo", "Descripción", "Obligatorio"],
    [
        ["id",                               "uuid",    "PK del cliente",                          "Sí"],
        ["name / commercial_name / legal_name","text",  "Nombres del establecimiento",             "Nullable"],
        ["email",                            "text",    "Email de contacto principal",              "Nullable"],
        ["nif",                              "text",    "NIF/CIF fiscal",                          "Nullable"],
        ["plan_id",                          "text",    "Plan contratado (FK → plans)",             "Nullable"],
        ["is_active",                        "bool",    "Estado activo/inactivo",                  "Sí (def. true)"],
        ["auth_user_id",                     "uuid",    "FK → Supabase Auth users",                "Nullable"],
        ["stripe_customer_id",               "text",    "ID en Stripe para facturación",           "Nullable"],
        ["trial_used",                       "bool",    "Si ya consumió el período de prueba",     "Nullable"],
        ["billing_frequency",                "text",    "monthly / annual",                        "Nullable"],
        ["iban / bank_name / swift",         "text",    "Datos bancarios para SEPA",              "Nullable"],
    ],
    col_widths=[2.5, 1.0, 2.8, 1.5]
)

add_heading(doc, "debacu_eval_guest_index — índice global de riesgo", 3, '2E75B6')
add_table(doc,
    ["Campo", "Tipo", "Descripción"],
    [
        ["identity_key",       "text",    "PK — hash criptográfico irreversible del huésped"],
        ["risk_band",          "text",    "HIGH / MEDIUM / LOW / NONE"],
        ["incidents_count",    "int",     "Total de incidencias acumuladas globalmente"],
        ["total_net_loss",     "numeric", "Pérdida económica neta acumulada (euros)"],
        ["stays_count",        "int",     "Total de estancias registradas en el sistema"],
        ["first_seen_date",    "date",    "Primera vez que aparece en el sistema"],
        ["last_incident_date", "date",    "Fecha de la incidencia más reciente"],
        ["doc_key",            "text",    "Hash del número de documento"],
        ["email_key",          "text",    "Hash del email"],
        ["phone_key",          "text",    "Hash del teléfono"],
    ],
    col_widths=[2.2, 1.0, 4.2]
)

add_heading(doc, "debacu_eval_identity_risk_state — estado calculado de riesgo", 3, '2E75B6')
add_table(doc,
    ["Campo", "Tipo", "Descripción"],
    [
        ["identity_key",             "text",    "PK — hash del huésped"],
        ["risk_level",               "enum",    "NONE / LOW / MEDIUM / HIGH / CRITICAL"],
        ["risk_score",               "numeric", "Score de 0 a 100"],
        ["incidents_total",          "int",     "Total de incidencias"],
        ["incidents_high",           "int",     "Incidencias de nivel alto"],
        ["incidents_critical",       "int",     "Incidencias de nivel crítico"],
        ["distinct_orgs_count",      "int",     "Número de organizaciones que han reportado"],
        ["distinct_properties_count","int",     "Número de propiedades distintas"],
        ["first_seen_at",            "timestamp","Cuándo se registró por primera vez"],
        ["last_incident_at",         "timestamp","Fecha del último incidente"],
        ["snapshot",                 "jsonb",   "Snapshot completo del estado para debugging"],
    ],
    col_widths=[2.5, 1.0, 3.9]
)

add_heading(doc, "debacu_eval_reservation_identities — datos encriptados", 3, '2E75B6')
add_table(doc,
    ["Campo", "Tipo", "Descripción"],
    [
        ["identity_key",       "text", "Hash derivado de los identificadores"],
        ["document_hash",      "text", "Hash SHA-256 del número de documento"],
        ["email_hash",         "text", "Hash SHA-256 del email"],
        ["phone_hash",         "text", "Hash SHA-256 del teléfono"],
        ["first_name_enc",     "text", "Nombre encriptado (AES, reversible con clave)"],
        ["last_name_enc",      "text", "Apellido encriptado"],
        ["identity_strength",  "text", "strong / medium / weak (cuántos IDs tiene)"],
        ["country",            "text", "País del documento"],
    ],
    col_widths=[2.0, 1.0, 4.4]
)

add_heading(doc, "3.3 Mapa de relaciones", 2, '2E75B6')
add_body(doc, (
    "customers (1)\n"
    "  ├── (1:1) debacu_eval_hotel_profile\n"
    "  ├── (1:1) debacu_eval_customer_profile\n"
    "  ├── (1:N) debacu_eval_invoices\n"
    "  └── (1:1) debacu_eval_customer_org_map\n"
    "              └── (1:1) debacu_eval_organizations\n"
    "                          ├── (1:N) debacu_eval_org_members\n"
    "                          ├── (1:N) debacu_eval_properties\n"
    "                          │           ├── (1:N) debacu_eval_property_room_types\n"
    "                          │           ├── (1:N) debacu_eval_property_seasons\n"
    "                          │           └── (1:N) debacu_eval_inventory_daily\n"
    "                          ├── (1:N) debacu_eval_import_batches\n"
    "                          │           └── (1:N) debacu_eval_import_rows\n"
    "                          ├── (1:N) debacu_eval_manual_checks\n"
    "                          │           └── (1:1) debacu_eval_manual_check_results\n"
    "                          ├── (1:N) debacu_eval_manual_incidents\n"
    "                          ├── (1:N) debacu_eval_reservations\n"
    "                          │           └── (1:N) debacu_eval_reservation_snapshots\n"
    "                          └── (1:N) pms_connections\n"
    "                                      └── (1:N) pms_sync_jobs\n\n"
    "identity_key (clave de dominio — no FK)\n"
    "  ├── debacu_eval_guest_index (índice global)\n"
    "  ├── debacu_eval_identity_risk_state (score actual)\n"
    "  ├── debacu_eval_identity_risk_events (log de cambios)\n"
    "  ├── debacu_eval_reservation_identities (datos encriptados)\n"
    "  └── debacu_eval_manual_incidents / manual_check_results"
), italic=True, color_hex='2E4057')

add_heading(doc, "3.4 Problemas detectados en el modelo de datos", 2, '2E75B6')
add_table(doc,
    ["Problema", "Tablas afectadas", "Impacto", "Prioridad"],
    [
        ["Tres variantes del índice de huéspedes sin sincronización garantizada",
         "guest_index / import_guest_index / org_guest_index", "ALTO — riesgo de inconsistencia de scores", "ALTA"],
        ["identity_key sin FK referencial (clave de dominio)",
         "Todas las tablas de riesgo", "MEDIO — huérfanos posibles si falla el pepper", "MEDIA"],
        ["Doble sistema de exportaciones sin migración completa",
         "audit_exports / debacu_eval_audit_exports / customer_audit_exports", "MEDIO — confusión en queries admin", "MEDIA"],
        ["Tablas de backup con fechas en el schema público",
         "debacu_evaluations_backup_20260207, guest_index_bak_20260314", "BAJO — espacio + confusión", "BAJA"],
        ["debacu_evaluations + manual_incidents coexisten sin unificación clara",
         "debacu_evaluations, debacu_eval_manual_incidents", "ALTO — deuda técnica visible", "ALTA"],
        ["customers.plan_id + subscriptions duplican el plan activo",
         "customers, subscriptions", "MEDIO — desincronización si Stripe webhook falla", "MEDIA"],
        ["Campo snapshot: jsonb sin schema TypeScript definido",
         "debacu_eval_identity_risk_state", "BAJO — dificulta debugging", "BAJA"],
    ],
    col_widths=[2.5, 2.2, 1.8, 1.0]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. ANÁLISIS DE NECESIDADES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. ANÁLISIS DE NECESIDADES", 1)

add_heading(doc, "4.1 Necesidades críticas — afectan a la retención y activación", 2, '2E75B6')
add_table(doc,
    ["Necesidad", "Problema actual", "Impacto en el negocio"],
    [
        ["Onboarding guiado en primer login",
         "El usuario llega al dashboard vacío sin saber qué hacer. El campo profile_completed existe pero no hay wizard forzado.",
         "Alta tasa de abandono en los primeros 7 días. Clients activados tardíamente o no activados."],
        ["Notificaciones proactivas de riesgo",
         "Cuando una reserva tiene un huésped con riesgo HIGH/CRITICAL, no se genera ningún aviso. El cliente debe consultar manualmente.",
         "El valor core del producto (protección ante riesgo) no se percibe en tiempo real. Diferencial perdido."],
        ["Watchlist con UI",
         "La tabla watchlist_reservations existe en BD pero no hay pantalla ni funcionalidad expuesta.",
         "Funcionalidad de alto valor sin monetizar. Los clientes no saben que pueden vigilar huéspedes específicos."],
        ["Explicación contextual del score de riesgo",
         "El resultado de búsqueda muestra el nivel (HIGH) pero no explica por qué: cuántas incidencias, en qué tipo, cuánta pérdida.",
         "El cliente no puede tomar una decisión informada sin más contexto."],
        ["Chatbot / asistente virtual",
         "No existe. Las dudas sobre privacidad, planes y uso generan tickets de soporte evitables.",
         "Coste de soporte alto. Fricción en onboarding. Oportunidad de upsell perdida."],
    ],
    col_widths=[1.8, 3.0, 2.6]
)

add_heading(doc, "4.2 Necesidades de producto — mejoran competitividad", 2, '2E75B6')
add_table(doc,
    ["Necesidad", "Estado actual", "Impacto"],
    [
        ["API pública documentada",               "api_token existe en customers, sin docs ni endpoints públicos",   "Abre mercado B2B, integraciones con channel managers y CRMs"],
        ["Benchmark sectorial en la UI",          "debacu_adr_reference_by_category existe en BD, sin UI",          "Diferenciador enorme: 'tu ADR está 8% por debajo de tu categoría'"],
        ["Dashboard de impacto económico",        "No existe",                                                       "El cliente ve cuánto dinero ha evitado perder gracias al sistema"],
        ["Resolución de incidencias",             "No hay flujo de 'incidencia resuelta' con rebaje de score",      "El score de un huésped nunca baja aunque haya pagado el daño"],
        ["Widget de consulta embebible",          "No existe",                                                       "Permite consulta de riesgo desde formularios de check-in propios"],
        ["Comparativa multi-propiedad",           "Solo para ENTERPRISE, sin UI específica",                        "Cadenas hoteleras ven su portfolio completo en un vistazo"],
        ["Alertas predictivas de revenue",        "Las alertas son reactivas, no predictivas",                      "'Tu pickup para el próximo puente está un 15% por debajo del año pasado'"],
        ["Módulo para arrendamiento",             "No existe, pero el modelo de datos es compatible",               "Amplía el mercado a gestores inmobiliarios y propietarios de alquiler"],
    ],
    col_widths=[2.0, 2.5, 3.0]
)

add_heading(doc, "4.3 Necesidades de UX/UI", 2, '2E75B6')
add_table(doc,
    ["Área", "Problema", "Mejora propuesta"],
    [
        ["Pantalla de búsqueda",       "Resultado sin contexto explicativo del riesgo",                      "Añadir explicación: '3 incidencias, 2 de nivel alto, pérdida acumulada €340'"],
        ["Pantalla de búsqueda",       "Sin botón 'Registrar incidencia sobre este huésped' en el resultado","Añadir CTA directo desde el resultado"],
        ["Pantalla de búsqueda",       "Sin historial de búsquedas recientes visible",                       "Panel lateral con últimas 10 consultas de la sesión"],
        ["Menú lateral",               "Usuario BASIC ve opciones PROFESSIONAL bloqueadas sin saber por qué","Tooltip explicativo + CTA de upgrade al hacer hover"],
        ["Mi cuenta",                  "No hay comparativa de planes dentro de la app",                      "Sección 'Ver todos los planes' accesible desde Mi cuenta"],
        ["Mi cuenta",                  "Sin aviso en-app cuando el trial está próximo a expirar",            "Banner de aviso 7 días antes + email automatizado"],
        ["Downgrade de plan",          "No se explica qué funcionalidades se perderán antes de confirmar",   "Modal de confirmación con lista clara de lo que se desactiva"],
        ["Solicitar acceso",           "Sin confirmación visual inmediata ni estimación de tiempo de revisión","Pantalla de confirmación: 'Tu solicitud se revisará en 24-48h'"],
        ["Revenue Intelligence",       "Demasiadas sub-rutas en el menú lateral para PROFESSIONAL",          "Agrupar en un único ítem expandible 'Revenue' con sub-ítems"],
        ["Móvil / tablet",             "UI no optimizada para pantallas pequeñas (recepción usa móvil)",     "Responsive design prioritario para pantallas de búsqueda y alertas"],
    ],
    col_widths=[1.5, 2.8, 3.1]
)

add_heading(doc, "4.4 Necesidades técnicas / deuda", 2, '2E75B6')
add_table(doc,
    ["Deuda técnica", "Impacto", "Esfuerzo de resolución"],
    [
        ["157 Edge Functions sin versionado ni agrupación formal",              "Mantenimiento costoso, difícil deprecar",  "ALTO (refactoring)"],
        ["Tres variantes del índice de huéspedes",                              "Riesgo de inconsistencia de datos",        "MEDIO"],
        ["debacu_evaluations coexistiendo con manual_incidents",                "Deuda técnica histórica visible",          "MEDIO"],
        ["Tablas backup en schema público",                                      "Confusión + espacio",                      "BAJO"],
        ["customers.plan_id duplicando subscriptions",                          "Desincronización posible",                 "BAJO"],
        ["Sin API Gateway — frontend llama directo a Edge Functions",           "Sin rate limiting, sin versionado de API", "ALTO"],
        ["Gestión de usuarios de una organización sin UI",                      "Solo 1 usuario por organización actualmente","MEDIO"],
        ["Sin ambiente de staging diferenciado del de producción",              "Riesgo en deploys",                        "MEDIO"],
    ],
    col_widths=[3.0, 2.2, 2.0]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PLAN DE ACCIÓN POR FASES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. PLAN DE ACCIÓN POR FASES", 1)

add_body(doc, (
    "El plan está diseñado para generar valor incremental desde la primera semana. "
    "Cada fase tiene criterios de éxito medibles y puede ejecutarse con un equipo pequeño. "
    "Las fases no son estrictamente secuenciales — algunas tareas de fases posteriores pueden "
    "adelantarse si los recursos lo permiten."
))

# Timeline visual
add_heading(doc, "5.0 Timeline de alto nivel", 2, '2E75B6')
add_table(doc,
    ["Fase", "Nombre", "Duración", "Foco principal", "Impacto esperado"],
    [
        ["FASE 1", "Activación y UX crítica",       "Semanas 1–4",  "Onboarding, notificaciones, UX búsqueda", "−30% abandono primer login, +NPS"],
        ["FASE 2", "Chatbot y soporte automático",  "Semanas 3–7",  "Chatbot FAQs + contextual",               "−40% tickets soporte"],
        ["FASE 3", "Funcionalidades de valor oculto","Semanas 5–10","Watchlist UI, dashboard impacto, API",    "+retención, +upsell ENTERPRISE"],
        ["FASE 4", "Diferenciación competitiva",    "Semanas 8–16", "Benchmark, alertas predictivas, módulos","Nuevos segmentos de mercado"],
    ],
    col_widths=[0.8, 2.5, 1.5, 2.5, 2.0]
)

doc.add_paragraph()

# FASE 1
add_heading(doc, "FASE 1 — Activación y UX crítica (Semanas 1–4)", 2, '1F5C99')
add_body(doc, "Objetivo: Que el cliente llegue, entienda el producto y complete su configuración en la primera sesión.")

add_heading(doc, "Tareas F1.1 — Onboarding guiado", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Tabla/Componente afectado", "Esfuerzo"],
    [
        ["T1.1.1", "Wizard de configuración inicial forzado en primer login si profile_completed = false. Pasos: perfil del establecimiento → primera consulta de prueba → tour del panel", "debacu_eval_hotel_profile, HotelProfileWizardDialog", "M"],
        ["T1.1.2", "Banner persistente en el dashboard si profile_completed = false con CTA al wizard", "AppShell, EvalAuthContext", "S"],
        ["T1.1.3", "Email automatizado de bienvenida con los 3 primeros pasos a seguir, enviado 1h después del primer login", "debacu_eval_email_log, Edge Function", "S"],
        ["T1.1.4", "Tooltip en menú lateral para ítems bloqueados por plan: 'Disponible en plan PROFESSIONAL — haz upgrade aquí'", "NavItem, AppShell", "S"],
    ],
    col_widths=[0.8, 3.5, 2.0, 0.8]
)

add_heading(doc, "Tareas F1.2 — Notificaciones proactivas de riesgo", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Tabla/Edge Function", "Esfuerzo"],
    [
        ["T1.2.1", "Crear sistema de notificaciones en-app: tabla notifications con org_id, type, message, read, created_at", "Nueva tabla: debacu_eval_notifications", "M"],
        ["T1.2.2", "Edge Function nocturna que cruza reservas próximas (check-in en ≤48h) con identity_risk_state y crea notificaciones para riesgo HIGH/CRITICAL", "debacu_eval_reservations × identity_risk_state", "M"],
        ["T1.2.3", "Badge de notificaciones en la barra superior del panel con acceso rápido a la lista", "AppShell, nuevo componente NotificationBell", "S"],
        ["T1.2.4", "Email diario de alertas: 'Tienes X huéspedes con riesgo alto que llegan mañana'", "debacu_eval_email_log, Brevo", "M"],
    ],
    col_widths=[0.8, 3.5, 2.2, 0.6]
)

add_heading(doc, "Tareas F1.3 — Mejoras UX pantalla de búsqueda", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T1.3.1", "Añadir en el resultado de búsqueda: número de incidencias, tipos más frecuentes, pérdida económica total, número de establecimientos que reportaron", "S"],
        ["T1.3.2", "Botón 'Registrar incidencia sobre este huésped' directamente visible en el resultado de búsqueda", "S"],
        ["T1.3.3", "Panel de historial de últimas 10 búsquedas de la sesión actual (visible en sidebar derecho)", "S"],
        ["T1.3.4", "Confirmación de solicitud de acceso: pantalla de éxito con estimación de tiempo '24–48h'", "XS"],
    ],
    col_widths=[0.8, 5.0, 0.8]
)

add_heading(doc, "KPIs de éxito Fase 1", 3, '2E75B6')
add_bullet(doc, "Tasa de completado de perfil en primer login: objetivo ≥ 70% (vs baseline a medir)")
add_bullet(doc, "Tasa de abandono en primeros 7 días: objetivo −30% vs baseline")
add_bullet(doc, "NPS del producto tras 30 días de uso: objetivo ≥ 7")
add_bullet(doc, "Reducción de tickets 'no sé cómo empezar': objetivo −50%")

doc.add_paragraph()

# FASE 2
add_heading(doc, "FASE 2 — Chatbot y soporte automático (Semanas 3–7)", 2, '1A7A3E')
add_body(doc, "Objetivo: Reducir la carga de soporte y aumentar la autonomía del usuario con un asistente conversacional.")

add_heading(doc, "Tareas F2.1 — Chatbot Fase 1 (FAQs estáticas)", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T2.1.1", "Crear base de conocimiento RAG en Supabase: tabla chat_kb_chunks con pgvector", "M"],
        ["T2.1.2", "Indexar todos los fragmentos de FAQ, guía de uso y niveles de riesgo (185+ fragmentos)", "M"],
        ["T2.1.3", "Edge Function chatbot_query: recibe mensaje → RAG → Claude → respuesta en JSON", "M"],
        ["T2.1.4", "Componente React ChatWidget: burbuja flotante, interfaz de chat, minimizable", "M"],
        ["T2.1.5", "Integrar en todas las rutas /app/* y en /planes, /solicitar-acceso", "S"],
        ["T2.1.6", "Logging en chat_sessions y chat_messages", "S"],
        ["T2.1.7", "Botón de escalado a soporte siempre visible", "XS"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "Tareas F2.2 — Chatbot Fase 2 (contextual)", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T2.2.1", "Inyectar contexto del usuario en el system prompt: plan, pantalla actual, profile_completed", "S"],
        ["T2.2.2", "Flujo de onboarding guiado vía chatbot: activación cuando profile_completed = false", "M"],
        ["T2.2.3", "Intención 'interpretar_riesgo': el chatbot explica el último resultado de búsqueda del usuario", "M"],
        ["T2.2.4", "Intención 'upgrade_plan': el chatbot guía al usuario al proceso de cambio de plan", "S"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "KPIs de éxito Fase 2", 3, '2E75B6')
add_bullet(doc, "Reducción de tickets de soporte: objetivo −40% en 60 días")
add_bullet(doc, "Tasa de resolución autónoma del chatbot: objetivo ≥ 75%")
add_bullet(doc, "NPS del chatbot: objetivo ≥ 7/10")

doc.add_paragraph()

# FASE 3
add_heading(doc, "FASE 3 — Funcionalidades de valor oculto (Semanas 5–10)", 2, '7B3F00')
add_body(doc, "Objetivo: Exponer valor que ya existe en el sistema pero que los clientes no pueden usar por falta de UI.")

add_heading(doc, "Tareas F3.1 — Watchlist de huéspedes", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T3.1.1", "Pantalla 'Watchlist': listado de huéspedes vigilados por la organización (usa tabla watchlist_reservations)", "M"],
        ["T3.1.2", "Botón 'Añadir a watchlist' en el resultado de búsqueda y en el detalle de incidencia", "S"],
        ["T3.1.3", "Integrar watchlist en el sistema de notificaciones (F1.2): si huésped en watchlist hace reserva → alerta inmediata", "M"],
        ["T3.1.4", "Exportación de la watchlist en PDF/CSV", "S"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "Tareas F3.2 — Dashboard de impacto económico", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T3.2.1", "Nueva sección 'Mi impacto' en el dashboard: 'Has consultado X huéspedes. Y tenían riesgo ALTO o CRÍTICO. Pérdida potencial evitada: €Z'", "M"],
        ["T3.2.2", "Cálculo del 'coste evitado': reservas de huéspedes HIGH/CRITICAL que el cliente consultó previamente × ADR × noches", "M"],
        ["T3.2.3", "Informe mensual de impacto por email (complementa el reporte semanal actual)", "S"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "Tareas F3.3 — Resolución de incidencias", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T3.3.1", "Añadir campo 'resolved_at', 'resolution_type', 'recovered_amount' en debacu_eval_manual_incidents", "S"],
        ["T3.3.2", "UI para marcar una incidencia como resuelta con tipo: 'pagó el daño', 'acuerdo extrajudicial', 'no recuperable'", "M"],
        ["T3.3.3", "Recalcular risk_score cuando se resuelve una incidencia (el score baja proporcionalmente)", "M"],
        ["T3.3.4", "Evento en identity_risk_events de tipo 'INCIDENT_RESOLVED'", "S"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "Tareas F3.4 — API pública (v1 básica)", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T3.4.1", "Documentar y habilitar los endpoints básicos de la API existente usando el api_token de customers", "M"],
        ["T3.4.2", "Endpoints v1: GET /risk/{identifier}, POST /incidents, GET /incidents (con paginación y filtros)", "M"],
        ["T3.4.3", "Portal de documentación de la API (Swagger/OpenAPI) accesible desde Mi cuenta", "M"],
        ["T3.4.4", "Rate limiting por api_token (máx. 1000 req/día en BASIC, 10.000 en ENTERPRISE)", "S"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "KPIs de éxito Fase 3", 3, '2E75B6')
add_bullet(doc, "Watchlist usada por ≥ 40% de los clientes activos en el primer mes")
add_bullet(doc, "NPS del panel: mejora de +1 punto vs Fase 1")
add_bullet(doc, "≥ 10 clientes usando la API en los primeros 60 días")

doc.add_paragraph()

# FASE 4
add_heading(doc, "FASE 4 — Diferenciación competitiva (Semanas 8–16+)", 2, '5C0099')
add_body(doc, "Objetivo: Convertir Debacu en el producto de referencia del sector con funcionalidades que ningún competidor tiene.")

add_heading(doc, "Tareas F4.1 — Benchmark sectorial", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T4.1.1", "Activar la tabla debacu_adr_reference_by_category en la UI del hotel profile y en el dashboard de Revenue", "S"],
        ["T4.1.2", "Widget 'Tu ADR vs sector': comparativa de tu ADR real vs el ADR de referencia de tu categoría en tu ciudad", "M"],
        ["T4.1.3", "Activar vista v_outlier_hotels para identificar outliers y mostrar al admin, potencialmente al cliente", "M"],
        ["T4.1.4", "Dashboard de benchmarking sectorial completo: RevPAR, ocupación, ADR vs categoría y geografía", "L"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "Tareas F4.2 — Alertas predictivas de revenue", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T4.2.1", "Modelo de predicción de pickup: comparar pace actual vs mismo período del año anterior → score de desviación", "L"],
        ["T4.2.2", "Alerta automática: 'Tu pickup para el próximo [período] está un X% por debajo del año pasado. Considera [acción]'", "M"],
        ["T4.2.3", "Sugerencias de precio basadas en pickup, temporada y eventos (usa debacu_eval_property_seasons)", "L"],
        ["T4.2.4", "Integrar alertas predictivas en el sistema de notificaciones (F1.2)", "S"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "Tareas F4.3 — Expansión a arrendamiento y otros sectores", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T4.3.1", "Adaptar el catálogo de incidencias para arrendamiento: impago de renta, daños en el inmueble, ocupación ilegal", "M"],
        ["T4.3.2", "Nuevo tipo de establecimiento en la solicitud de acceso: 'gestora inmobiliaria', 'propietario de alquiler'", "S"],
        ["T4.3.3", "Plan específico para arrendamiento con pricing diferenciado", "M"],
        ["T4.3.4", "Módulo de seguimiento de arrendatarios con historial de pagos y contratos", "L"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "Tareas F4.4 — Gestión multi-usuario por organización", 3, '2E75B6')
add_table(doc,
    ["Tarea", "Descripción", "Esfuerzo"],
    [
        ["T4.4.1", "UI de gestión de miembros de la organización: invitar usuario, asignar rol (admin/operador/solo lectura)", "M"],
        ["T4.4.2", "Flujo de invitación por email similar al de activación de cuenta", "S"],
        ["T4.4.3", "Permisos por rol: solo el admin puede registrar incidencias, el operador solo puede consultar", "M"],
        ["T4.4.4", "Log de auditoría diferenciado por usuario dentro de la misma organización", "S"],
    ],
    col_widths=[0.8, 4.8, 0.8]
)

add_heading(doc, "KPIs de éxito Fase 4", 3, '2E75B6')
add_bullet(doc, "Benchmark sectorial: feature más valorada en NPS survey")
add_bullet(doc, "Alertas predictivas: ≥ 60% de clientes con Revenue Intelligence las tienen activadas")
add_bullet(doc, "Módulo arrendamiento: 50+ clientes en 6 meses de disponibilidad")
add_bullet(doc, "Multi-usuario: ≥ 30% de organizaciones con 2+ usuarios activos")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. RESUMEN DE PRIORIDADES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. MATRIZ DE PRIORIDADES — IMPACTO VS ESFUERZO", 1)

add_table(doc,
    ["Prioridad", "Tarea", "Fase", "Impacto", "Esfuerzo", "Quick win"],
    [
        ["1",  "Notificaciones proactivas de riesgo (F1.2)",                 "1", "CRÍTICO",   "M",  "Sí"],
        ["2",  "Wizard de onboarding en primer login (T1.1.1)",              "1", "MUY ALTO",  "M",  "No"],
        ["3",  "Chatbot FAQs básico (F2.1)",                                 "2", "ALTO",      "M",  "No"],
        ["4",  "Resultado de búsqueda con contexto explicativo (T1.3.1)",    "1", "ALTO",      "S",  "Sí"],
        ["5",  "Botón 'Registrar incidencia' desde resultado (T1.3.2)",      "1", "ALTO",      "S",  "Sí"],
        ["6",  "Tooltip en ítems bloqueados por plan (T1.1.4)",              "1", "MEDIO",     "S",  "Sí"],
        ["7",  "Watchlist UI (F3.1)",                                        "3", "ALTO",      "M",  "No"],
        ["8",  "Dashboard de impacto económico (F3.2)",                      "3", "MUY ALTO",  "M",  "No"],
        ["9",  "Resolución de incidencias (F3.3)",                           "3", "ALTO",      "M",  "No"],
        ["10", "Benchmark sectorial ADR (T4.1.1–T4.1.2)",                   "4", "ALTO",      "S",  "Sí"],
        ["11", "API pública v1 (F3.4)",                                      "3", "MEDIO",     "M",  "No"],
        ["12", "Alertas predictivas de revenue (F4.2)",                      "4", "MUY ALTO",  "L",  "No"],
        ["13", "Multi-usuario por organización (F4.4)",                      "4", "MEDIO",     "M",  "No"],
        ["14", "Módulo arrendamiento (F4.3)",                                "4", "ALTO",      "L",  "No"],
        ["15", "Limpieza de deuda técnica BD (tablas duplicadas/backup)",    "1", "BAJO",      "S",  "Sí"],
    ],
    col_widths=[0.6, 3.5, 0.7, 1.1, 0.9, 1.0]
)

add_heading(doc, "Leyenda de esfuerzo", 2, '2E75B6')
add_body(doc, "XS: < 1 día  |  S: 1–3 días  |  M: 1–2 semanas  |  L: 2–4 semanas  |  XL: > 1 mes", italic=True)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. PUNTOS FUERTES Y DIFERENCIADORES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. PUNTOS FUERTES Y DIFERENCIADORES ACTUALES", 1)

add_table(doc,
    ["Fortaleza", "Detalle", "Por qué es diferencial"],
    [
        ["Privacidad by design real",
         "Hashing irreversible con pepper global. Encriptación de nombres. Sin PII en logs ni resultados.",
         "Competidores tienen listas negras identificables. Debacu puede operar en toda la UE sin fricción legal."],
        ["Auditoría granular",
         "Cada consulta, cada exportación queda registrada con SHA-256, destinatario y base legal.",
         "Valor probatorio en reclamaciones judiciales. Diferencial para clientes que operan en sectores regulados."],
        ["Revenue + Riesgo en el mismo producto",
         "No es solo riesgo. La combinación con revenue intelligence hace el producto sticky.",
         "Los pure players de riesgo no tienen Revenue. Los de Revenue no tienen riesgo. Debacu tiene ambos."],
        ["Multi-tenancy sólido",
         "RLS a nivel de base de datos. Imposible acceder a datos de otra organización.",
         "Confianza real del cliente en la privacidad de sus datos frente a competidores."],
        ["Integración PMS (ENTERPRISE)",
         "Sincronización automática con los principales PMS del mercado.",
         "Elimina la fricción de la importación manual. Retención alta en clientes ENTERPRISE."],
        ["Modelo de planes escalable",
         "BASIC → PROFESSIONAL → ENTERPRISE con upsell natural bien definido.",
         "Adquisición en el plan BASIC, monetización vía upgrade orgánico."],
    ],
    col_widths=[1.8, 2.8, 3.0]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. RIESGOS DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. RIESGOS Y MITIGACIONES", 1)

add_table(doc,
    ["Riesgo", "Probabilidad", "Impacto", "Mitigación"],
    [
        ["Rechazo de clientes al sistema por percepción de fichero de personas",
         "MEDIA", "ALTO",
         "FAQs de privacidad muy visibles. Chatbot con respuestas claras sobre hashing. Comunicación proactiva en onboarding."],
        ["Deuda técnica (157 Edge Functions) ralentiza nuevas features",
         "ALTA", "MEDIO",
         "Priorizar refactoring incremental. Cada nueva feature evita crear Edge Functions nuevas donde sea posible."],
        ["Desincronización de plan_id entre customers y subscriptions si Stripe falla",
         "BAJA", "ALTO",
         "Añadir job de reconciliación diaria. Alertar al admin si hay discrepancia."],
        ["Entrada de competidor con modelo similar y más recursos",
         "MEDIA", "ALTO",
         "Acelerar las fases 3 y 4. El benchmark sectorial y la API pública son los mejores fosos defensivos."],
        ["Reclamación RGPD por parte de un huésped",
         "BAJA", "MUY ALTO",
         "Proceso documentado de derecho de supresión. DPA en regla. Datos en hashes irreversibles (base legal sólida)."],
        ["Crecimiento de BD sin índices adecuados afecta tiempos de consulta",
         "MEDIA", "MEDIO",
         "Revisar índices en identity_risk_state, guest_index y reservations. La tabla risk_alerts tiene índices definidos en migración."],
    ],
    col_widths=[2.5, 1.2, 1.0, 3.0]
)

# ── Guardar
output_path = r"c:\DebacuEvaluation\DEBACU_Producto_Estado_y_Plan_de_Accion.docx"
doc.save(output_path)
print(f"Documento guardado: {output_path}")
