"""
Añade sección de guía de colores al documento DEBACU_Chatbot_Implantacion.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tblPr.append(borders)

def add_heading(doc, text, level=1, color_hex='1F3864'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    run.font.size = Pt(sizes.get(level, 11))
    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text, bold=False, italic=False, color_hex=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    if color_hex:
        r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)

def add_table(doc, headers, rows, col_widths=None, header_bg='1F3864'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F2F5FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_color_swatch(doc, hex_fill, label, description):
    """Añade una fila de color como párrafo con bloque visual."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    run_label = p.add_run(f"  {label}  ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    r, g, b = tuple(int(hex_fill[i:i+2], 16) for i in (0, 2, 4))
    run_label.font.color.rgb = RGBColor(r, g, b)
    run_desc = p.add_run(f" — {description}")
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = RGBColor(60, 60, 60)

# ── Carga el documento existente ─────────────────────────────────────────────
path = r"c:\DebacuEvaluation\DEBACU_Chatbot_Implantacion.docx"
doc = Document(path)

# ── Añade salto de página antes de la nueva sección ──────────────────────────
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# NUEVA SECCIÓN: GUÍA DE COLORES E INDICADORES VISUALES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. GUÍA DE COLORES E INDICADORES VISUALES", 1)
add_body(doc, (
    "Esta sección documenta el significado de cada color y elemento visual en el panel de Debacu Evaluation. "
    "Es fundamental para el chatbot y para el equipo de soporte, ya que los usuarios frecuentemente preguntan "
    "'¿qué significa este color?' o '¿qué quiere decir el punto rojo?'"
))

# 8.1 Colores de nivel de riesgo
add_heading(doc, "8.1 Colores de nivel de riesgo (badges y etiquetas)", 2, '2E75B6')
add_body(doc,
    "Los badges de riesgo aparecen en: resultados de búsqueda, panel de alertas, screening CSV, "
    "historial de búsquedas recientes y cualquier lugar donde se muestre el nivel de riesgo de un huésped.",
    italic=True
)

add_table(doc,
    ["Color visual", "Nivel", "Código interno", "Significado y acción recomendada"],
    [
        ["🔴 Rojo (bg-red)",
         "ALTO",
         "HIGH",
         "Historial serio con pérdidas significativas en otros establecimientos. Se recomienda solicitar depósito máximo, pedir garantías adicionales o valorar denegar el servicio."],
        ["🟠 Ámbar/Naranja (bg-amber)",
         "MEDIO",
         "MEDIUM",
         "Varios incidentes o impacto económico moderado. Tomar precauciones: pedir depósito, revisar historial propio, estar alerta durante la estancia."],
        ["🟢 Verde (bg-emerald/green)",
         "BAJO",
         "LOW",
         "Algún incidente aislado de poco impacto. Precaución mínima, registro habitual. Sin señales graves."],
        ["⚫ Gris oscuro (bg-slate)",
         "SIN SEÑALES",
         "NONE",
         "El huésped no tiene historial registrado en la plataforma. No significa que sea una persona 'buena', simplemente no hay incidencias previas."],
        ["⬛ Gris claro (bg-slate-100)",
         "NO CONCLUYENTE",
         "NO_CONCLUYENTE",
         "La búsqueda no ha podido determinar una identidad única (por ejemplo, búsqueda por nombre completo sin documento). El resultado no es fiable con ese criterio."],
    ],
    col_widths=[1.5, 1.0, 1.5, 3.4]
)

# 8.2 Colores del score numérico
add_heading(doc, "8.2 Colores del score numérico (barra de progreso)", 2, '2E75B6')
add_body(doc,
    "El score Debacu es un número del 0 al 100. La barra de progreso cambia de color según el rango:",
    italic=True
)
add_color_swatch(doc, 'DC2626', "ROJO — Score 60–100", "Riesgo alto. Múltiples incidencias graves o pérdidas significativas.")
add_color_swatch(doc, 'D97706', "ÁMBAR — Score 30–59", "Riesgo medio. Incidencias moderadas o impacto económico notable.")
add_color_swatch(doc, '10B981', "VERDE — Score 0–29", "Riesgo bajo o sin señales relevantes.")
doc.add_paragraph()
add_body(doc,
    "El panel de score también cambia su fondo: rojo claro para alto, ámbar claro para medio, verde claro para bajo.",
    italic=True, color_hex='666666'
)

# 8.3 Puntos de color en historial reciente
add_heading(doc, "8.3 Puntos de color en el historial de búsquedas recientes", 2, '2E75B6')
add_body(doc,
    "En el panel de búsquedas recientes (visible entre el formulario de búsqueda y los resultados), "
    "cada búsqueda pasada muestra un pequeño punto de color que indica el nivel de riesgo obtenido en esa consulta:",
    italic=True
)
add_color_swatch(doc, 'F87171', "Punto ROJO",    "Esa búsqueda devolvió nivel de riesgo ALTO.")
add_color_swatch(doc, 'FBBF24', "Punto ÁMBAR",   "Esa búsqueda devolvió nivel de riesgo MEDIO.")
add_color_swatch(doc, '34D399', "Punto VERDE",   "Esa búsqueda devolvió nivel de riesgo BAJO.")
add_color_swatch(doc, 'CBD5E1', "Punto GRIS",    "Sin señales, no concluyente, o búsqueda sin resultado.")
doc.add_paragraph()

# 8.4 Panel explicativo del riesgo
add_heading(doc, "8.4 Panel explicativo '¿Por qué este nivel de riesgo?'", 2, '2E75B6')
add_body(doc,
    "Cuando una búsqueda GLOBAL devuelve señales, aparece un panel explicativo debajo del resultado principal. "
    "El fondo de ese panel también es coloreado:",
    italic=True
)
add_color_swatch(doc, 'FEE2E2', "Fondo ROJO CLARO",   "El nivel de riesgo es ALTO. El panel explica cuántas incidencias graves hay y la pérdida acumulada.")
add_color_swatch(doc, 'FFFBEB', "Fondo ÁMBAR CLARO",  "El nivel de riesgo es MEDIO.")
add_color_swatch(doc, 'ECFDF5', "Fondo VERDE CLARO",  "El nivel de riesgo es BAJO.")
doc.add_paragraph()

# 8.5 Colores de severidad en incidencias propias
add_heading(doc, "8.5 Colores de severidad en incidencias propias (modo 'Mis registros')", 2, '2E75B6')
add_body(doc,
    "En el modo 'Mis registros', cada incidencia muestra un badge de severidad. "
    "La severidad es distinta al nivel de riesgo: indica cómo de grave fue esa incidencia concreta.",
    italic=True
)
add_table(doc,
    ["Color", "Severidad", "Código", "Descripción"],
    [
        ["🟢 Verde",   "Baja",    "LOW",      "Incidencia menor o de impacto muy limitado."],
        ["🟠 Ámbar",   "Media",   "MEDIUM",   "Incidencia notable. Requiere atención."],
        ["🔴 Rojo",    "Alta",    "HIGH",     "Incidencia grave. Pérdida o daño significativo."],
        ["🔴 Rojo",    "Crítica", "CRITICAL", "Incidencia muy grave. Acción inmediata recomendada."],
    ],
    col_widths=[1.2, 1.2, 1.2, 3.8]
)

# 8.6 Colores de estado en Screening CSV
add_heading(doc, "8.6 Colores en el Screening CSV (importación masiva)", 2, '2E75B6')
add_body(doc,
    "Cuando se procesa un archivo CSV con una lista de huéspedes, los resultados se clasifican por risk_band:",
    italic=True
)
add_table(doc,
    ["Color", "Risk Band", "Significado"],
    [
        ["🔴 Rojo",   "HIGH",   "Huésped con historial grave. Revisar antes del check-in."],
        ["🟠 Ámbar",  "MEDIUM", "Huésped con historial moderado. Tomar precauciones."],
        ["🟢 Verde",  "LOW",    "Huésped con historial leve o sin incidencias relevantes."],
        ["⚫ Gris",   "NONE",   "Sin historial en la plataforma."],
    ],
    col_widths=[1.2, 1.2, 5.0]
)

# 8.7 FAQs sobre colores (para el chatbot)
add_heading(doc, "8.7 FAQs sobre colores — respuestas para el chatbot", 2, '2E75B6')
add_body(doc, "Estas respuestas deben indexarse en la base de conocimiento RAG del chatbot:", bold=True)

color_faqs = [
    ("¿Qué significa el color rojo en el resultado de búsqueda?",
     "El color rojo indica que el huésped tiene un nivel de riesgo ALTO. Tiene historial serio de incidencias en otros establecimientos, con pérdidas económicas significativas. Se recomienda solicitar garantías máximas o valorar denegar el servicio."),
    ("¿Qué significa el color naranja o ámbar?",
     "El color ámbar indica nivel de riesgo MEDIO. El huésped tiene varios incidentes registrados o el impacto económico es moderado. Toma precauciones adicionales como pedir un depósito."),
    ("¿Qué significa el color verde en el riesgo?",
     "Verde indica nivel de riesgo BAJO. Puede haber algún incidente aislado pero de poco impacto. Procede con precaución normal."),
    ("¿Qué significa que no haya color o que sea gris?",
     "Gris o sin color significa que el huésped no tiene historial registrado en la plataforma (nivel NONE). Esto no confirma que sea una persona sin incidencias, simplemente que no hay registros previos en Debacu Evaluation."),
    ("¿Qué es el punto de color que aparece junto a las búsquedas recientes?",
     "Es un indicador rápido del resultado de esa búsqueda anterior. Punto rojo = riesgo alto en esa consulta. Punto ámbar = riesgo medio. Punto verde = riesgo bajo. Punto gris = sin señales o resultado no concluyente. Puedes hacer clic en la búsqueda para repetirla."),
    ("¿Qué significa el panel de fondo rojo que aparece en el resultado?",
     "Es el panel explicativo '¿Por qué este nivel de riesgo?'. El fondo rojo confirma que el nivel es ALTO y dentro encontrarás el detalle: cuántas incidencias graves tiene, en cuántos establecimientos ha sido reportado y cuánta pérdida económica acumulada tiene."),
    ("¿La barra de progreso del score qué indica?",
     "La barra de progreso muestra visualmente el score de riesgo de 0 a 100. Si está en rojo y llena, el score es alto (60–100). Si está en ámbar, el score es medio (30–59). Si está en verde, el score es bajo (0–29). A mayor barra y más roja, mayor es el riesgo histórico de ese huésped."),
    ("¿El color rojo significa que debo denegar la reserva?",
     "No necesariamente. El color rojo (riesgo ALTO) es una señal de alerta, no una orden. La decisión final siempre es tuya. Lo que el sistema te recomienda es: solicitar garantías máximas, pedir depósito, o revisar el detalle de las incidencias. En casos de riesgo CRÍTICO (score muy alto y múltiples incidencias graves), la recomendación es denegar el servicio, pero siempre queda a tu criterio."),
    ("¿Por qué el mismo huésped puede aparecer con distintos colores en distintas consultas?",
     "El nivel de riesgo se actualiza continuamente. Si otros establecimientos han registrado nuevas incidencias desde tu última consulta, el score puede haber subido. Por eso es importante consultar justo antes del check-in y no fiarse de una consulta de semanas atrás."),
    ("¿Los colores de severidad en mis registros son lo mismo que el nivel de riesgo global?",
     "No. La severidad (en 'Mis registros') indica cómo de grave fue esa incidencia concreta que tú registraste: LOW, MEDIUM, HIGH, CRITICAL. El nivel de riesgo global (NONE/LOW/MEDIUM/HIGH en el modo Global) es un cálculo agregado de todas las incidencias de todos los establecimientos para esa identidad."),
]

for q, a in color_faqs:
    add_body(doc, f"P: {q}", bold=True)
    add_body(doc, f"R: {a}")
    doc.add_paragraph()

# 8.8 Tabla resumen para soporte rápido
add_heading(doc, "8.8 Tabla resumen — referencia rápida de colores", 2, '2E75B6')
add_table(doc,
    ["Color", "Dónde aparece", "Qué significa", "Acción sugerida al usuario"],
    [
        ["🔴 Rojo",      "Badge de riesgo, barra score, fondo panel, punto historial",       "Riesgo ALTO / Severidad HIGH-CRITICAL / Score 60-100", "Garantías máximas o valorar denegar"],
        ["🟠 Ámbar",     "Badge de riesgo, barra score, fondo panel, punto historial",       "Riesgo MEDIO / Severidad MEDIUM / Score 30-59",        "Pedir depósito, estar alerta"],
        ["🟢 Verde",     "Badge de riesgo, barra score, fondo panel, punto historial",       "Riesgo BAJO / Severidad LOW / Score 0-29",             "Precaución normal"],
        ["⚫ Gris oscuro","Badge de riesgo, punto historial",                                 "SIN SEÑALES (NONE)",                                    "Proceder con normalidad"],
        ["⬛ Gris claro", "Badge de riesgo",                                                  "NO CONCLUYENTE",                                        "Buscar con otro dato (doc, email, tel)"],
    ],
    col_widths=[1.0, 2.5, 2.5, 1.8]
)

# ── Guarda ────────────────────────────────────────────────────────────────────
doc.save(path)
print(f"Sección de colores añadida: {path}")
